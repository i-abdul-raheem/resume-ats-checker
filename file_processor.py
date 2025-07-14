import os
import logging
from typing import Optional
from pathlib import Path
import tempfile

# Import file processing libraries
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. DOCX files will not be supported.")

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PyPDF2 not available. PDF files will not be supported.")

logger = logging.getLogger(__name__)

class FileProcessor:
    """Utility class for processing different file formats and extracting text."""
    
    def __init__(self):
        self.supported_formats = []
        
        if DOCX_AVAILABLE:
            self.supported_formats.append('.docx')
        if PDF_AVAILABLE:
            self.supported_formats.append('.pdf')
    
    def extract_text_from_file(self, file_path: str) -> Optional[str]:
        """
        Extract text content from a file based on its extension.
        
        Args:
            file_path: Path to the file to process
            
        Returns:
            Extracted text content or None if processing failed
        """
        try:
            file_path = Path(file_path)
            file_extension = file_path.suffix.lower()
            
            if not file_path.exists():
                logger.error(f"File not found: {file_path}")
                return None
            
            if file_extension == '.docx' and DOCX_AVAILABLE:
                return self._extract_from_docx(file_path)
            elif file_extension == '.pdf' and PDF_AVAILABLE:
                return self._extract_from_pdf(file_path)
            elif file_extension == '.txt':
                return self._extract_from_txt(file_path)
            else:
                logger.error(f"Unsupported file format: {file_extension}")
                return None
                
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            return None
    
    def extract_text_from_upload(self, file_content: bytes, filename: str) -> Optional[str]:
        """
        Extract text content from uploaded file content.
        
        Args:
            file_content: Raw file content as bytes
            filename: Original filename with extension
            
        Returns:
            Extracted text content or None if processing failed
        """
        try:
            # Create a temporary file to process
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            # Extract text from the temporary file
            text_content = self.extract_text_from_file(temp_file_path)
            
            # Clean up temporary file
            os.unlink(temp_file_path)
            
            return text_content
            
        except Exception as e:
            logger.error(f"Error processing uploaded file {filename}: {e}")
            return None
    
    def _extract_from_docx(self, file_path: Path) -> Optional[str]:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            return '\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX file {file_path}: {e}")
            return None
    
    def _extract_from_pdf(self, file_path: Path) -> Optional[str]:
        """Extract text from PDF file."""
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(page_text.strip())
            
            return '\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF file {file_path}: {e}")
            return None
    
    def _extract_from_txt(self, file_path: Path) -> Optional[str]:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding if UTF-8 fails
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Error reading TXT file {file_path}: {e}")
                return None
        except Exception as e:
            logger.error(f"Error reading TXT file {file_path}: {e}")
            return None
    
    def get_supported_formats(self) -> list:
        """Get list of supported file formats."""
        return self.supported_formats.copy()
    
    def is_supported_format(self, filename: str) -> bool:
        """Check if the file format is supported."""
        file_extension = Path(filename).suffix.lower()
        return file_extension in self.supported_formats or file_extension == '.txt' 